#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
将 docs/操作手册 下的 Markdown 手册分别导出为 Word 文档。

导出规则：
1. 每个 Markdown 生成一个同名 docx 文件。
2. 标题、正文、列表、步骤、表格、图片按较接近原文的方式落入 Word。
3. 图片优先直接嵌入文档，便于转发。
"""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Inches


ROOT_DIR = Path(__file__).resolve().parents[1]
MANUAL_DIR = ROOT_DIR / "docs" / "操作手册"
OUTPUT_DIR = MANUAL_DIR / "word"

IMAGE_PATTERN = re.compile(r"!\[(.*?)\]\((.*?)\)")
HEADING_PATTERN = re.compile(r"^(#{1,6})\s+(.*)")
BULLET_PATTERN = re.compile(r"^[-*]\s+(.*)")
NUMBER_PATTERN = re.compile(r"^(\d+)\.\s+(.*)")


def iter_manual_files() -> list[Path]:
    """按文件名顺序返回全部操作手册 Markdown 文件。"""

    return sorted(MANUAL_DIR.glob("*.md"), key=lambda item: item.name)


def add_paragraph(document: Document, text: str, style: str | None = None) -> None:
    """在 Word 中追加段落。"""

    paragraph = document.add_paragraph(style=style)
    paragraph.add_run(text)


def flush_table(document: Document, table_lines: list[str]) -> None:
    """将连续的 Markdown 表格行写入 Word 表格。"""

    if not table_lines:
        return

    parsed_rows: list[list[str]] = []
    for line in table_lines:
        cells = [cell.strip() for cell in line.strip().strip("|").split("|")]
        if not cells:
            continue
        if all(re.fullmatch(r"[-:\s]+", cell or "") for cell in cells):
            continue
        parsed_rows.append(cells)

    if not parsed_rows:
        return

    col_count = max(len(row) for row in parsed_rows)
    table = document.add_table(rows=0, cols=col_count)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for row_index, row_values in enumerate(parsed_rows):
        row_cells = table.add_row().cells
        for col_index in range(col_count):
            row_cells[col_index].text = row_values[col_index] if col_index < len(row_values) else ""
        if row_index == 0:
            for cell in row_cells:
                for run in cell.paragraphs[0].runs:
                    run.bold = True

    document.add_paragraph("")


def export_one_markdown(markdown_file: Path) -> Path:
    """导出单个 Markdown 文件到 Word。"""

    document = Document()
    table_buffer: list[str] = []

    for raw_line in markdown_file.read_text(encoding="utf-8").splitlines():
        line = raw_line.rstrip()
        stripped = line.strip()

        if not stripped:
            flush_table(document, table_buffer)
            table_buffer.clear()
            continue

        if stripped.startswith("|") and stripped.endswith("|"):
            table_buffer.append(stripped)
            continue

        flush_table(document, table_buffer)
        table_buffer.clear()

        image_match = IMAGE_PATTERN.fullmatch(stripped)
        if image_match:
            alt_text, image_path = image_match.groups()
            image_file = (markdown_file.parent / image_path).resolve()
            if alt_text:
                add_paragraph(document, alt_text)
            if image_file.exists():
                document.add_picture(str(image_file), width=Inches(5.8))
            else:
                add_paragraph(document, f"图片缺失：{image_path}")
            document.add_paragraph("")
            continue

        heading_match = HEADING_PATTERN.match(stripped)
        if heading_match:
            hashes, title = heading_match.groups()
            document.add_heading(title, level=min(len(hashes), 6))
            continue

        bullet_match = BULLET_PATTERN.match(stripped)
        if bullet_match:
            add_paragraph(document, bullet_match.group(1), style="List Bullet")
            continue

        number_match = NUMBER_PATTERN.match(stripped)
        if number_match:
            index_text, content = number_match.groups()
            add_paragraph(document, f"{index_text}. {content}", style="List Number")
            continue

        add_paragraph(document, stripped)

    flush_table(document, table_buffer)

    output_file = OUTPUT_DIR / f"{markdown_file.stem}.docx"
    document.save(output_file)
    return output_file


def main() -> None:
    """批量导出全部操作手册。"""

    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    exported_files = [export_one_markdown(markdown_file) for markdown_file in iter_manual_files()]

    print("exported_count=", len(exported_files))
    for file in exported_files:
        print(file.relative_to(ROOT_DIR))


if __name__ == "__main__":
    main()
